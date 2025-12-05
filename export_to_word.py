from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import os
from datetime import datetime

# Путь к папке с вашими файлами
#folder_path = "C:/Users/Nurba/PycharmProjects/service_with_mini_app"  # на работе папка
folder_path = "C:/Users/User/PycharmProjects/test"  # домашняя папка

# Получаем текущую дату в формате дд.мм
current_date = datetime.now().strftime("%d.%m")
output_file = f"all_project_files_{current_date}.docx"  # Имя выходного файла с датой

# Список допустимых расширений файлов
ALLOWED_EXTENSIONS = {'.html', '.css', '.py', '.yaml', '.yml'}
ALLOWED_FILENAMES = {'Dockerfile'}  # Точное имя файла без расширения

# Создаём новый документ Word
doc = Document()

# Добавляем стиль "Code", если его нет
styles = doc.styles
if "Code" not in styles:
    style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Courier New"  # Моноширинный шрифт для кода
    style.font.size = Pt(10)  # Размер шрифта 10pt
    style.paragraph_format.space_after = Pt(6)  # Отступ после абзаца 6pt

# Добавляем заголовок документа
doc.add_heading("Список файлов проекта", level=1)

# Проходим по всем файлам в папке и её подпапках
for root, dirs, files in os.walk(folder_path):
    # Пропускаем папки __pycache__ и venv
    if "__pycache__" in dirs:
        dirs.remove("__pycache__")
    if "venv" in dirs:
        dirs.remove("venv")

    for file in files:
        # Проверяем расширение файла или точное имя
        file_extension = os.path.splitext(file)[1].lower()  # Получаем расширение файла
        file_name = file  # Полное имя файла

        if file_extension in ALLOWED_EXTENSIONS or file_name in ALLOWED_FILENAMES:
            file_path = os.path.join(root, file)
            doc.add_heading(f"Файл: {file_path}", level=2)

            # Читаем содержимое файла
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                    doc.add_paragraph(content, style="Code")  # Добавляем содержимое с форматированием
            except Exception as e:
                doc.add_paragraph(f"Ошибка чтения файла {file_path}: {e}", style="Normal")

            doc.add_page_break()  # Разделяем файлы новой страницей

# Сохраняем документ
doc.save(output_file)
print(f"Файлы сохранены в {output_file}")